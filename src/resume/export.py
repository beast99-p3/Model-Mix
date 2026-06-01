from __future__ import annotations

import re
from pathlib import Path

from src.resume.format import (
    bullet_parts,
    is_blank_line,
    is_contact_line,
    is_section_heading,
    line_role,
    normalize_text_chars,
    split_lines,
)

SKILL_CATEGORY_RE = re.compile(r"^([^:]{3,40}:)\s*(.+)$")


def _add_bottom_border(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _set_run_font(run, *, size_pt: float, bold: bool = False, underline: bool = False) -> None:
    from docx.shared import Pt

    run.bold = bold
    run.underline = underline
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)


def _replace_paragraph_text(paragraph, text: str, *, bold_prefix: bool = False) -> None:
    text = normalize_text_chars(text.strip())
    if paragraph.runs:
        for run in paragraph.runs[1:]:
            run.text = ""
        first = paragraph.runs[0]
    else:
        first = paragraph.add_run("")
    if bold_prefix:
        m = SKILL_CATEGORY_RE.match(text)
        if m:
            first.text = m.group(1) + " "
            _set_run_font(first, size_pt=10.5, bold=True)
            if len(paragraph.runs) > 1:
                second = paragraph.runs[1]
            else:
                second = paragraph.add_run("")
            second.text = m.group(2)
            _set_run_font(second, size_pt=10.5)
            return
    first.text = text
    _set_run_font(first, size_pt=10.5)


def _indent_from_source(source_line: str) -> float:
    from docx.shared import Inches

    spaces = len(source_line) - len(source_line.lstrip())
    return Inches(min(spaces, 24) / 16)


def write_docx_from_source_template(
    source_path: Path,
    docx_path: Path,
    draft_text: str,
    source_text: str,
) -> Path | None:
    """Replace text in the user's uploaded .docx, keeping all original styles."""
    if source_path.suffix.lower() != ".docx" or not source_path.is_file():
        return None

    import docx
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    draft_lines = split_lines(normalize_text_chars(draft_text))
    source_lines = split_lines(normalize_text_chars(source_text))
    line_idx = 0

    def next_line() -> str | None:
        nonlocal line_idx
        while line_idx < len(draft_lines):
            current = draft_lines[line_idx]
            line_idx += 1
            return current
        return None

    doc = docx.Document(str(source_path))
    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            para = Paragraph(child, doc)
            if not para.text.strip():
                continue
            new_line = next_line()
            if new_line is None:
                break
            _, content = bullet_parts(new_line)
            _replace_paragraph_text(
                para,
                content if bullet_parts(new_line)[0].strip() else new_line.strip(),
                bold_prefix=bool(bullet_parts(source_lines[min(line_idx - 1, len(source_lines) - 1)])[0].strip())
                if line_idx > 0
                else False,
            )
        elif tag == "tbl":
            table = Table(child, doc)
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if not para.text.strip():
                            continue
                        new_line = next_line()
                        if new_line is None:
                            break
                        _, content = bullet_parts(new_line)
                        _replace_paragraph_text(
                            para,
                            content if bullet_parts(new_line)[0].strip() else new_line.strip(),
                        )

    doc.save(docx_path)
    return docx_path


def write_docx_from_text(docx_path: Path, draft_text: str, source_text: str) -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    draft_lines = split_lines(normalize_text_chars(draft_text))
    source_lines = split_lines(normalize_text_chars(source_text))
    if len(source_lines) != len(draft_lines):
        n = max(len(source_lines), len(draft_lines))
        source_lines += [""] * (n - len(source_lines))
        draft_lines += [""] * (n - len(draft_lines))

    doc = Document()
    for i, (src, out) in enumerate(zip(source_lines, draft_lines)):
        role = line_role(source_lines, i)
        text = normalize_text_chars(out if out.strip() else src).rstrip()

        if role == "blank" or not text:
            doc.add_paragraph("")
            continue

        if role == "name":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text.strip())
            _set_run_font(run, size_pt=14, bold=True)
            continue

        if role == "contact":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text.strip())
            _set_run_font(run, size_pt=10)
            continue

        if role == "heading":
            p = doc.add_paragraph()
            run = p.add_run(text.strip())
            _set_run_font(run, size_pt=11, bold=True, underline=True)
            _add_bottom_border(p)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            continue

        prefix, content = bullet_parts(out)
        if role == "bullet" and content:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = _indent_from_source(src) or Inches(0.25)
            p.paragraph_format.space_after = Pt(1)
            _replace_paragraph_text(p, content, bold_prefix=True)
            continue

        if role == "continuation":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = _indent_from_source(src)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text.strip())
            _set_run_font(run, size_pt=10.5)
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text.strip())
        _set_run_font(run, size_pt=10.5)

    doc.save(docx_path)
    return docx_path


def write_pdf_from_text(pdf_path: Path, draft_text: str, source_text: str) -> Path:
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    draft_lines = split_lines(normalize_text_chars(draft_text))
    source_lines = split_lines(normalize_text_chars(source_text))
    if len(source_lines) != len(draft_lines):
        n = max(len(source_lines), len(draft_lines))
        source_lines += [""] * (n - len(source_lines))
        draft_lines += [""] * (n - len(draft_lines))

    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    name_style = ParagraphStyle("Name", parent=styles["Normal"], fontName="Times-Bold", fontSize=14, alignment=TA_CENTER, spaceAfter=4)
    contact_style = ParagraphStyle("Contact", parent=styles["Normal"], fontName="Times-Roman", fontSize=10, alignment=TA_CENTER, spaceAfter=6)
    heading_style = ParagraphStyle("Heading", parent=styles["Normal"], fontName="Times-Bold", fontSize=11, alignment=TA_LEFT, spaceBefore=6, spaceAfter=2)
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontName="Times-Roman", fontSize=10.5, leading=12, alignment=TA_LEFT, spaceAfter=1)
    cont_style = ParagraphStyle("Cont", parent=body_style, leftIndent=24, spaceAfter=0)
    bullet_style = ParagraphStyle("Bullet", parent=body_style, leftIndent=18, bulletIndent=0, spaceAfter=1)

    def esc(s: str) -> str:
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def skill_markup(s: str) -> str:
        m = SKILL_CATEGORY_RE.match(s.strip())
        if m:
            return f"<b>{esc(m.group(1))}</b> {esc(m.group(2))}"
        return esc(s)

    story: list = []
    for i, (src, out) in enumerate(zip(source_lines, draft_lines)):
        role = line_role(source_lines, i)
        text = normalize_text_chars(out if out.strip() else src).strip()
        if role == "blank" or not text:
            story.append(Spacer(1, 5))
            continue
        if role == "name":
            story.append(Paragraph(esc(text), name_style))
        elif role == "contact":
            story.append(Paragraph(esc(text), contact_style))
        elif role == "heading":
            story.append(Paragraph(f"<u><b>{esc(text)}</b></u>", heading_style))
        elif role == "bullet":
            _, content = bullet_parts(out)
            story.append(Paragraph(f"&bull; {skill_markup(content or text)}", bullet_style))
        elif role == "continuation":
            story.append(Paragraph(esc(text), cont_style))
        else:
            story.append(Paragraph(esc(text), body_style))

    doc.build(story)
    return pdf_path
